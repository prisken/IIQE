#!/usr/bin/env python3
"""Extract study manuals, high-density guides, and question banks from content source DOCX only."""
from __future__ import annotations

import json
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

CONTENT = Path("/Users/priskenlo/IIQE papers/content source")
OUT = Path(__file__).resolve().parents[1] / "public" / "data"

PAPERS = {
    1: {
        "id": 1,
        "code": "P1",
        "titleZh": "保險原理及實務",
        "titleEn": "Principles and Practice of Insurance",
        "manual": "Paper 1 - Insurance Principles and Practice (Student Rewrite).docx",
        "guide": "Paper 1 - High-Density Study Guide (Deck Rewrite).docx",
        "questions": "Paper 1 - Original Bilingual Mock Questions.docx",
        "exam": {"count": 75, "minutes": 120, "passPercent": 70},
        "weights": [
            {"id": "1", "weight": 12, "titleZh": "風險及保險", "titleEn": "Risk and Insurance"},
            {"id": "2", "weight": 16, "titleZh": "法律原則", "titleEn": "Legal Principles"},
            {"id": "3", "weight": 30, "titleZh": "保險原則", "titleEn": "Principles of Insurance"},
            {"id": "4", "weight": 9, "titleZh": "保險公司的主要功能", "titleEn": "Core Functions of an Insurance Company"},
            {"id": "5", "weight": 5, "titleZh": "香港保險業的結構", "titleEn": "Structure of the HK Insurance Industry"},
            {"id": "6", "weight": 21, "titleZh": "保險業的規管架構", "titleEn": "Regulatory Framework"},
            {"id": "7", "weight": 7, "titleZh": "職業道德及其他有關問題", "titleEn": "Ethical and Other Issues"},
        ],
    },
    2: {
        "id": 2,
        "code": "P2",
        "titleZh": "一般保險",
        "titleEn": "General Insurance",
        "manual": "Paper 2 - General Insurance (Student Rewrite).docx",
        "guide": "Paper 2 - High-Density Study Guide (Deck Rewrite).docx",
        "questions": "Paper 2 - Original Bilingual Mock Questions.docx",
        "exam": {"count": 50, "minutes": 75, "passPercent": 70},
        "weights": [
            {"id": "1", "weight": 46, "titleZh": "保險產品", "titleEn": "Insurance Products"},
            {"id": "2", "weight": 34, "titleZh": "核保及保單措辭", "titleEn": "Underwriting and Policy Wording"},
            {"id": "3", "weight": 10, "titleZh": "索償", "titleEn": "Claims"},
            {"id": "4", "weight": 10, "titleZh": "客戶服務", "titleEn": "Customer Service"},
        ],
    },
    3: {
        "id": 3,
        "code": "P3",
        "titleZh": "長期保險",
        "titleEn": "Long Term Insurance",
        "manual": "Paper 3 - Long Term Insurance (Student Rewrite).docx",
        "guide": "Paper 3 - High-Density Study Guide (Deck Rewrite).docx",
        "questions": "Paper 3 - Original Bilingual Mock Questions.docx",
        "exam": {"count": 50, "minutes": 75, "passPercent": 70},
        "weights": [
            {"id": "1", "weight": 10, "titleZh": "人壽保險簡介", "titleEn": "Introduction to Life Insurance"},
            {"id": "2", "weight": 20, "titleZh": "人壽保險產品", "titleEn": "Life Insurance Products"},
            {"id": "3", "weight": 24, "titleZh": "人壽保險附約及其他產品", "titleEn": "Riders and Other Products"},
            {"id": "4", "weight": 24, "titleZh": "人壽保險的運作", "titleEn": "Operation of Life Insurance"},
            {"id": "5", "weight": 22, "titleZh": "銷售過程及客戶保障", "titleEn": "Sales Process and Customer Protection"},
        ],
    },
    4: {
        "id": 4,
        "code": "P4",
        "titleZh": "強制性公積金計劃",
        "titleEn": "Mandatory Provident Fund Schemes",
        "manual": "Paper 4 - MPF Schemes (Student Rewrite).docx",
        "guide": "Paper 4 - High-Density Study Guide (Deck Rewrite).docx",
        "questions": "Paper 4 - Original Bilingual Mock Questions.docx",
        "exam": {"count": 80, "minutes": 120, "passPercent": 70},
        "weights": [
            {"id": "1", "weight": 1, "titleZh": "強積金制度簡介", "titleEn": "MPF System Introduction"},
            {"id": "2", "weight": 6, "titleZh": "規管架構", "titleEn": "Regulatory Framework"},
            {"id": "3", "weight": 45, "titleZh": "強積金制度主要特點", "titleEn": "Main Features of MPF System"},
            {"id": "4", "weight": 5, "titleZh": "強積金受託人", "titleEn": "MPF Trustees"},
            {"id": "5", "weight": 19, "titleZh": "強積金計劃及投資", "titleEn": "MPF Schemes and Investment"},
            {"id": "6", "weight": 4, "titleZh": "職業退休計劃與強積金銜接", "titleEn": "ORSO–MPF Interface"},
            {"id": "7", "weight": 20, "titleZh": "強積金中介人", "titleEn": "MPF Intermediaries"},
        ],
    },
    5: {
        "id": 5,
        "code": "P5",
        "titleZh": "投資相連長期保險",
        "titleEn": "Investment-linked Long Term Insurance",
        "manual": "Paper 5 - Investment-linked Long Term Insurance (Student Rewrite).docx",
        "guide": "Paper 5 - High-Density Study Guide (Deck Rewrite).docx",
        "questions": "Paper 5 - Original Bilingual Mock Questions.docx",
        "exam": {"count": 80, "minutes": 120, "passPercent": 70},
        "weights": [
            {"id": "1", "weight": 2.5, "titleZh": "投資相連壽險簡介", "titleEn": "Introduction to ILAS"},
            {"id": "2", "weight": 20, "titleZh": "投資概念", "titleEn": "Investment Concepts"},
            {"id": "3", "weight": 35, "titleZh": "投資工具及基金", "titleEn": "Investment Vehicles and Funds"},
            {"id": "4", "weight": 32.5, "titleZh": "投資相連壽險產品", "titleEn": "ILAS Products"},
            {"id": "5", "weight": 10, "titleZh": "規管及銷售", "titleEn": "Regulation and Sales"},
        ],
    },
}

REF_RE = re.compile(
    r"^(?P<ref>\d+(?:\.\d+)*[a-z]?(?:\([a-z]\))?|\d+\([a-z]\))\s*[　\s]+(?P<title>.+)$",
    re.I,
)
CN_NUM = {
    "一": "1",
    "二": "2",
    "三": "3",
    "四": "4",
    "五": "5",
    "六": "6",
    "七": "7",
    "八": "8",
    "九": "9",
    "十": "10",
}


def parse_chapter_heading(text: str) -> tuple[str, str] | None:
    """Return (chapterId, title) for chapter headings across paper formats."""
    t = text.strip()
    t = re.sub(r"^[（(]續[）)]\s*", "", t)
    t = re.sub(r"[（(]續[）)]\s*$", "", t).strip()

    m = re.match(r"^Chapter\s*(?P<n>\d+)\s*(?P<title>.*)$", t, re.I)
    if m:
        return m.group("n"), (m.group("title") or "").strip("　 ").strip() or f"Chapter {m.group('n')}"

    m = re.match(r"^第\s*(?P<n>\d+)\s*章\s*(?P<title>.*)$", t)
    if m:
        return m.group("n"), (m.group("title") or "").strip("　 ").strip() or f"第{m.group('n')}章"

    m = re.match(r"^第\s*(?P<cn>[一二三四五六七八九十]+)\s*章\s*(?P<title>.*)$", t)
    if m and m.group("cn") in CN_NUM:
        n = CN_NUM[m.group("cn")]
        return n, (m.group("title") or "").strip("　 ").strip() or f"第{n}章"

    m = re.match(r"^(?P<n>\d+)\s+(?P<title>.+)$", t)
    if m and not re.match(r"^\d+\.\d+", t):
        # Avoid treating "1.1 title" as chapter; bare "1 Title" is chapter
        return m.group("n"), m.group("title").strip("　 ").strip()

    return None


def is_chapter_start(text: str, level: int | None, kind: str) -> bool:
    if parse_chapter_heading(text):
        if kind == "guide" and level in (1, 2, None):
            return True
        if kind == "manual" and level == 1:
            return True
    return False


def iter_block_items(parent):
    from docx.oxml.ns import qn

    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def para_text(p: Paragraph) -> str:
    return (p.text or "").replace("\xa0", " ").strip()


def table_markdown(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    cols = rows[0].count("|") - 1
    sep = "| " + " | ".join(["---"] * max(cols, 1)) + " |"
    return "\n".join([rows[0], sep, *rows[1:]])


def heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    m = re.match(r"Heading\s*(\d+)", style_name, re.I)
    return int(m.group(1)) if m else None


def normalize_ref(raw: str) -> str:
    raw = raw.strip()
    raw = re.sub(r"^Ch\s*\d+\s+", "", raw, flags=re.I)
    return raw


def chapter_of_ref(ref: str) -> str:
    ref = normalize_ref(ref)
    m = re.match(r"^(\d+)", ref)
    if m:
        return m.group(1)
    if "操守" in ref:
        return "7"
    if "附錄" in ref or "附录" in ref:
        return "3"
    return "0"


def extract_study_doc(path: Path, kind: str) -> dict:
    """Return {chapters: [{id, title, sections: [{id, title, blocks: [{type, text/html}]}]}]}"""
    doc = Document(str(path))
    chapters: list[dict] = []
    current_ch = None
    current_sec = None
    started = False

    def ensure_chapter(cid: str, title: str):
        nonlocal current_ch, current_sec
        for ch in chapters:
            if ch["id"] == cid:
                current_ch = ch
                current_sec = None
                return
        current_ch = {"id": cid, "title": title, "sections": []}
        chapters.append(current_ch)
        current_sec = None

    def ensure_section(sid: str, title: str):
        nonlocal current_sec
        if not current_ch:
            ensure_chapter(chapter_of_ref(sid) or "0", title)
        for sec in current_ch["sections"]:
            if sec["id"] == sid:
                current_sec = sec
                return
        current_sec = {"id": sid, "title": title, "blocks": []}
        current_ch["sections"].append(current_sec)

    def add_block(block: dict):
        nonlocal current_sec
        if not current_ch:
            return
        if not current_sec:
            ensure_section(f"{current_ch['id']}.0", current_ch["title"])
        current_sec["blocks"].append(block)

    for block in iter_block_items(doc):
        if isinstance(block, Table):
            if not started:
                continue
            md = table_markdown(block)
            if md:
                add_block({"type": "table", "markdown": md})
            continue

        text = para_text(block)
        if not text:
            continue
        level = heading_level(block.style.name if block.style else None)

        # Skip front matter / TOC until first real chapter heading
        if not started:
            if is_chapter_start(text, level, kind):
                # Avoid TOC: for manuals, require Heading 1; for guides Heading 2 "Chapter N" is fine
                if kind == "manual" and level != 1:
                    continue
                if kind == "manual" and text in ("目錄", "应考须知", "應考須知"):
                    continue
                started = True
            else:
                continue

        if text in ("目錄", "应考须知", "應考須知", "附件", "附錄") or text.startswith("附件"):
            # stop at appendices for cleaner navigation (keep content before)
            if current_ch and level == 1 and ("附件" in text or "附錄" in text or "附件" in text):
                break
            if level == 1 and text in ("目錄", "应考须知", "應考須知"):
                continue

        ch = parse_chapter_heading(text)
        if ch and ((kind == "manual" and level == 1) or (kind == "guide" and level in (1, 2))):
            ensure_chapter(ch[0], ch[1])
            continue

        if level in (2, 3, 4, 5) or (kind == "guide" and level in (3, 4)):
            m = REF_RE.match(text) or re.match(
                r"^(?P<ref>\d+(?:\.\d+)*[a-z]?(?:\([a-z]\))?)\s*(?P<title>.*)$", text
            )
            if m:
                ref = m.group("ref")
                title = (m.group("title") or "").strip() or ref
                # Skip bare section titles without numeric ref for guide H4 like "功能分類"
                if re.match(r"^\d+", ref):
                    ensure_section(ref, title)
                    continue
            # Guide H4 key points without numeric refs — attach under current section as titled blocks
            if kind == "guide" and level == 4 and current_ch:
                if not current_sec:
                    ensure_section(f"{current_ch['id']}.key", current_ch["title"])
                add_block({"type": "point", "title": text, "text": ""})
                continue

        if text.startswith("官方") or text.startswith("考試比重") or text.startswith("官方比重") or (
            "考試比重" in text[:12]
        ):
            if current_ch and not current_sec:
                ensure_section(f"{current_ch['id']}.meta", "考試比重")
            add_block({"type": "meta", "text": text})
            continue

        # Append body text to last point title if open empty point
        if current_sec and current_sec["blocks"]:
            last = current_sec["blocks"][-1]
            if last.get("type") == "point" and last.get("text") == "" and level is None:
                last["text"] = text
                continue

        add_block({"type": "p", "text": text})

    # Drop empty / meta-only junk sections with no content
    for ch in chapters:
        ch["sections"] = [s for s in ch["sections"] if s["blocks"] or s["id"].endswith(".meta") is False]
        for s in ch["sections"]:
            s["preview"] = next((b["text"] for b in s["blocks"] if b.get("type") == "p" and b.get("text")), "")
    return {"chapters": chapters}


OPTION_LINE_RE = re.compile(
    r"^(?P<letter>[A-D])(?P<sep>[.\u3001]\s*|\s+)(?P<rest>.+)$",
    re.S,
)
# English lines that begin with A/An/The… must not be treated as option markers.
ROMAN_LIST_RE = re.compile(
    r"^[ivxIVX]{1,4}\b|^[ivxIVX]{1,4}\s*,|^\(?[ivxIVX]{1,4}\)?\s*[,\u3001]",
)


def _looks_like_option_rest(letter: str, sep: str, rest: str) -> bool:
    """True if text after A-D is a real option body, not English prose starting with 'A '."""
    rest = rest.strip()
    if not rest:
        return False
    # Explicit delimiter: "A." / "A、" — always an option marker
    if "." in sep or "\u3001" in sep:
        return True
    # "A i, ii" / "B ii, iii, iv" style keys (lowercase, but real options)
    if ROMAN_LIST_RE.match(rest):
        return True
    # Only letter "A" is ambiguous with the English article:
    #   "A financial loss that will certainly occur"
    #   "A major earthquake affecting a whole community"
    # Sequential matching already blocks these when expected is B/C/D; this blocks
    # them if they ever appear when expected is still A (shouldn't for options).
    # Do NOT apply to B/C/D — legitimate options like "B nothing" / "D nothing".
    if letter == "A" and re.match(r"^[a-z]", rest):
        return False
    # Accepted for the expected letter:
    # - Chinese / 「引號」/ 《條例》
    # - Capitalized English ("Uncertainty…", "Nothing")
    # - Lowercase B/C/D bodies ("nothing")
    # - Numbers / parentheses
    return True


def _is_option_line(line: str, expected: str) -> re.Match[str] | None:
    """Accept an A-D line only if it is the next expected letter and looks like an option."""
    m = OPTION_LINE_RE.match(line)
    if not m:
        return None
    letter = m.group("letter")
    if letter != expected:
        # e.g. already on B; English "A financial loss..." must stay under B
        return None
    if not _looks_like_option_rest(letter, m.group("sep"), m.group("rest")):
        return None
    return m


def split_question_stem(stem: str) -> dict:
    """Parse bilingual stem + A-D options with sequential letter validation."""
    # Normalize DOCX newlines and HTML breaks into line list
    text = stem.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    parts = [p.strip() for p in text.split("\n") if p.strip()]

    expected = "A"
    options: list[dict] = []
    body_parts: list[str] = []

    for line in parts:
        m = _is_option_line(line, expected) if expected <= "D" else None
        if m:
            options.append({"letter": m.group("letter"), "text": m.group("rest").strip()})
            expected = chr(ord(expected) + 1) if expected < "D" else "E"
            continue
        if options:
            options[-1]["text"] += "\n" + line
        else:
            body_parts.append(line)

    # Fallback: same sequential rules if primary pass failed
    if len(options) != 4:
        options = []
        body_parts = []
        expected = "A"
        buf: list[str] = []
        for line in parts:
            m = _is_option_line(line, expected) if expected <= "D" else None
            if m:
                if not options:
                    body_parts = buf
                    buf = []
                options.append({"letter": m.group("letter"), "text": m.group("rest").strip()})
                expected = chr(ord(expected) + 1) if expected < "D" else "E"
            elif options:
                options[-1]["text"] += "\n" + line
            else:
                buf.append(line)
        if not options:
            body_parts = buf

    stem_text = "\n".join(body_parts) if body_parts else "\n".join(parts[:2])
    return {
        "stem": stem_text,
        "stemLines": body_parts,
        "options": options,
    }


def validate_question(q: dict) -> list[str]:
    errs = []
    letters = [o.get("letter") for o in q.get("options", [])]
    if letters != ["A", "B", "C", "D"]:
        errs.append(f"options letters={letters}")
    for o in q.get("options", []):
        if not (o.get("text") or "").strip():
            errs.append(f"empty option {o.get('letter')}")
    if q.get("answer") not in "ABCD":
        errs.append(f"bad answer {q.get('answer')!r}")
    elif q.get("answer") not in letters:
        errs.append(f"answer {q.get('answer')} missing from options")
    return errs


def extract_questions(path: Path) -> list[dict]:
    doc = Document(str(path))
    if not doc.tables:
        return []
    table = doc.tables[0]
    out = []
    warnings = []
    for i, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        if i == 0 or len(cells) < 4:
            continue
        if "題號" in cells[0]:
            continue
        m = re.match(r"^Q?(\d+)\s*/\s*(.+)$", cells[0].replace("\n", " "))
        if not m:
            continue
        qn = int(m.group(1))
        ref = normalize_ref(m.group(2))
        parsed = split_question_stem(cells[1])
        answer = cells[2].strip().upper()[:1]
        if answer not in "ABCD":
            warnings.append(f"Q{qn}: skip — answer {cells[2]!r}")
            continue
        item = {
            "id": qn,
            "ref": ref,
            "chapter": chapter_of_ref(ref),
            "stem": parsed["stem"],
            "stemLines": parsed["stemLines"],
            "options": parsed["options"],
            "answer": answer,
            "explanation": cells[3].strip(),
        }
        errs = validate_question(item)
        if errs:
            warnings.append(f"Q{qn}: {'; '.join(errs)}")
        out.append(item)
    if warnings:
        print(f"  question parse warnings ({len(warnings)}):")
        for w in warnings[:20]:
            print(f"    {w}")
        if len(warnings) > 20:
            print(f"    … +{len(warnings) - 20} more")
    return out


def write_json(path: Path, data):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _section_has_body(sec: dict) -> bool:
    for b in sec.get("blocks") or []:
        t = b.get("type")
        if t == "table" and (b.get("markdown") or "").strip():
            return True
        if t == "point" and ((b.get("title") or "").strip() or (b.get("text") or "").strip()):
            return True
        if t in ("p", "meta") and (b.get("text") or "").strip():
            return True
    return False


def _parent_section_id(sid: str) -> str | None:
    m = re.match(r"^(.*\.\d+)[a-z]+$", sid, re.I)
    if m:
        return m.group(1)
    if "." not in sid:
        return None
    return sid.rsplit(".", 1)[0]


def enrich_study_doc(doc: dict) -> None:
    """Mark heading-only vs content sections; drop empty trailing meta stubs."""
    for ch in doc.get("chapters") or []:
        sections = ch.get("sections") or []
        ids = [s["id"] for s in sections]
        cleaned = []
        for sec in sections:
            # Drop empty auto meta placeholders
            if sec["id"].endswith(".meta") and not _section_has_body(sec):
                continue
            kids = [x for x in ids if _parent_section_id(x) == sec["id"]]
            sec["hasContent"] = _section_has_body(sec)
            sec["isBranch"] = bool(kids)
            sec["childIds"] = kids
            if not sec["hasContent"] and not sec["isBranch"] and not sec.get("blocks"):
                continue
            cleaned.append(sec)
        ch["sections"] = cleaned


def _first_contentful_descendant(manual: dict, section_id: str, chapter_id: str | None) -> str | None:
    chapters = manual.get("chapters") or []
    pool = []
    for ch in chapters:
        if chapter_id and ch["id"] != chapter_id:
            continue
        pool.extend(ch.get("sections") or [])
    # Prefer exact contentful
    for sec in pool:
        if sec["id"] == section_id and sec.get("hasContent"):
            return sec["id"]
    # Then contentful descendants
    prefix = section_id + "."
    descendants = [s for s in pool if s["id"].startswith(prefix) and s.get("hasContent")]
    descendants.sort(key=lambda s: (s["id"].count("."), s["id"]))
    if descendants:
        return descendants[0]["id"]
    # Keep branch id if it exists (UI shows overview)
    for sec in pool:
        if sec["id"] == section_id:
            return sec["id"]
    return None


def link_guide_to_manual(guide: dict, manual: dict) -> None:
    """Point guide sections at the best readable manual section (avoid blank parents when possible)."""
    for ch in guide.get("chapters") or []:
        for sec in ch.get("sections") or []:
            candidates = [sec["id"]]
            parts = sec["id"].split(".")
            while len(parts) > 1:
                parts.pop()
                candidates.append(".".join(parts))
            target = None
            for c in candidates:
                hit = _first_contentful_descendant(manual, c, ch["id"])
                if hit:
                    target = hit
                    break
            if not target:
                for c in candidates:
                    hit = _first_contentful_descendant(manual, c, None)
                    if hit:
                        target = hit
                        break
            sec["manualTarget"] = target


def main():
    index = []
    for pid, meta in PAPERS.items():
        print(f"Extracting paper {pid}...")
        pdir = OUT / f"paper{pid}"
        manual = extract_study_doc(CONTENT / meta["manual"], "manual")
        guide = extract_study_doc(CONTENT / meta["guide"], "guide")
        questions = extract_questions(CONTENT / meta["questions"])

        enrich_study_doc(manual)
        enrich_study_doc(guide)
        link_guide_to_manual(guide, manual)

        by_chapter: dict[str, list] = defaultdict(list)
        for q in questions:
            by_chapter[q["chapter"]].append(q["id"])

        paper_meta = {
            "id": pid,
            "code": meta["code"],
            "titleZh": meta["titleZh"],
            "titleEn": meta["titleEn"],
            "exam": meta["exam"],
            "weights": meta["weights"],
            "stats": {
                "manualChapters": len(manual["chapters"]),
                "guideChapters": len(guide["chapters"]),
                "questions": len(questions),
                "questionsByChapter": {k: len(v) for k, v in sorted(by_chapter.items())},
            },
        }
        write_json(pdir / "meta.json", paper_meta)
        write_json(pdir / "manual.json", manual)
        write_json(pdir / "guide.json", guide)
        write_json(pdir / "questions.json", {"questions": questions})
        index.append(paper_meta)
        print(
            f"  manual ch={len(manual['chapters'])} guide ch={len(guide['chapters'])} "
            f"questions={len(questions)}"
        )

    write_json(OUT / "papers.json", {"papers": index})
    print(f"Done → {OUT}")


if __name__ == "__main__":
    main()
